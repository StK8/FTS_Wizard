import re
import time
from datetime import datetime
from docx import Document
import os
import glob
import math

from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.shared import Inches

from table_parse import table_parse
from table_parse import table_processing
from table_parse import generate_stations_text
from table_parse import generate_stations_summary_text
from image_export import image_export
from res_engineers import res_engineers
from fields import fields
from table_parse import JOB_TOOL_TYPES
from ttkbootstrap.dialogs import Messagebox

# # redirecting traceback during script execution to a file
# file = open("error_log.txt", "w")
# sys.stderr.write = file.write
#
# # if no XPT run value = '', else value = 'XPT-'. Used to generate docx report name and docx tags
# XPT_RUN = 'XPT-'
#
# IMAGE_WIDTH_VERTICAL = Inches(7.2)
# IMAGE_WIDTH_HORIZONTAL = Inches(10.5)
# DIRECTORY = os.getcwd()
# FILENAME_DOCX_TEMPLATE = 'MDT_template.docx'
# FILENAME_PPTX = 'BRRI_1101_0_8.5in_MDT_Sampling_Summary_9.pptx'
# FILENAME_DOCX_REPORT = re.search('.+(?=in_)', FILENAME_PPTX).group() + 'in_' + XPT_RUN + \
#                                     re.search('(?<=in_)(.+)(?=_Sampling)', FILENAME_PPTX).group() + '_Report'

def modify_docx(FILENAME_PPTX, FILENAME_DOCX_TEMPLATE, FILENAME_DOCX_REPORT, IMAGE_WIDTH_VERTICAL, IMAGE_WIDTH_HORIZONTAL, XPT, PRETEST_TOOLS, progress_bar, progress_label):


    exp_date = datetime(year=2026, month=2, day=20)
    now = datetime.now()

    if now > exp_date:
        return None

    # export images from ppt
    image_export(FILENAME_PPTX)

    progress_bar['value'] += 20
    progress_label['text'] = 'Images export from ppt...[OK]'
    progress_bar.update_idletasks()

    # generate DFA/sampling station text for report summary
    stations_summary_text = generate_stations_summary_text(table_processing(table_parse(FILENAME_PPTX)))

    progress_bar['value'] += 20
    progress_label['text'] = 'Generating stations summary text...[OK]'
    progress_bar.update_idletasks()

    #print(stations_summary_text)
    # generate text for each DFA/Sampling station
    stations_data = generate_stations_text(table_processing(table_parse(FILENAME_PPTX)))
    #print(stations_data)
    #print(JOB_TOOL_TYPES)

    progress_bar['value'] += 20
    progress_label['text'] = 'Generating text for each DFA station...[OK]'
    progress_bar.update_idletasks()

    document = Document(FILENAME_DOCX_TEMPLATE)

    # change docx 'Tags' property
    document.core_properties.keywords = FILENAME_DOCX_REPORT
    document.core_properties.title = FILENAME_DOCX_REPORT

    # for style in document.styles:
    #     print("style.name == %s" % style.name)

    # EDIT REPORT HEADER
    # enter tool type for at report header
    XPT_TEXT = 'PressureXpress'
    MDT_TEXT = 'Modular Formation Dynamics Tester'
    ORA_TEXT = 'ORA'
    Tool_type = ''

    WELL_NAME = re.search('.+(?=_\d+(\.){0,1}\d+in_)', FILENAME_PPTX).group()
    WELL_DIAMETER = re.search('(?<=_)(\d+(\.){0,1}\d+)(?=[Ii][Nn]_)', FILENAME_PPTX).group()

    # generate text for <TOOL_TYPE> placeholder in the report header
    if ('MDT' or 'Mdt') in FILENAME_PPTX and ('ORA' or 'Ora') in FILENAME_PPTX:
        Tool_type = MDT_TEXT + '\n'*4 + ORA_TEXT
    elif ('MDT' or 'Mdt') in FILENAME_PPTX and XPT == 'XPT-':
        Tool_type = XPT_TEXT + '\n'*4 + MDT_TEXT
    elif ('ORA' or 'Ora') in FILENAME_PPTX and XPT == 'XPT-':
        Tool_type = XPT_TEXT + '\n'*4 + ORA_TEXT
    elif ('MDT' or 'Mdt') in FILENAME_PPTX:
        Tool_type = MDT_TEXT
    elif ('ORA' or 'Ora') in FILENAME_PPTX:
        Tool_type = ORA_TEXT

    # generate text for <TOOL_TYPE> placeholder in the report header
    TOOLS_LIST = []
    if XPT == 'XPT-':
        TOOLS_LIST.append('XPT')
    if 'MDT' in FILENAME_PPTX:
        TOOLS_LIST.append('MDT')
    if 'ORA' in FILENAME_PPTX:
        TOOLS_LIST.append('ORA')



    # generate text for <WELL_NAME_DIAMETER> and <TOOLS-LIST> placeholders in the report header
    for p in document.paragraphs:
        for run in p.runs:
            if '<WELL_NAME_DIAMETER>' in run.text:
                run.text = WELL_NAME + '_' + WELL_DIAMETER + 'in'
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            if '<TOOLS-LIST>' in run.text:
                Tools_list = ''
                for tool in TOOLS_LIST:
                    Tools_list += tool + '-'
                Tools_list = Tools_list.strip("-")
                run.text = run.text.replace("<TOOLS-LIST>", Tools_list)
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            if '<PRETEST_TOOL>' in run.text:
                run.text = (PRETEST_TOOLS[0] + PRETEST_TOOLS[1] + PRETEST_TOOLS[2]).strip("-")
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    # insert tool type name, well name, well diameter, tools list, reservoir engineer
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '<TOOL_TYPE>' in run.text:
                            run.text = run.text.replace("<TOOL_TYPE>", Tool_type)
                            run.style = 'SLB_InSituPro_Doc_Cover_Font_Arial_18BW'
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        elif '<WELL_NAME>' in run.text:
                            run.text = run.text.replace("<WELL_NAME>", WELL_NAME)
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        elif '<FIELD_NAME>' in run.text:
                            field_name = re.search('[A-Z]{4}', WELL_NAME).group()
                            if field_name in fields:
                                run.text = run.text.replace("<FIELD_NAME>", fields.get(field_name))
                                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        elif '<WELL_DIAMETER>' in run.text:
                            run.text = run.text.replace("<WELL_DIAMETER>", WELL_DIAMETER + ' in')
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        elif '<TOOLS_LIST>' in run.text:
                            Tools_list = ''
                            for tool in TOOLS_LIST:
                                Tools_list += tool + ', '
                            Tools_list = Tools_list.strip(", ")
                            run.text = run.text.replace("<TOOLS_LIST>", Tools_list)
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        elif '<RES_ENG>' in run.text:
                            re_alias = os.getlogin()
                            re_name = res_engineers.get(re_alias)
                            run.text = run.text.replace("<RES_ENG>", re_name)
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        elif '<REPORT_DATE>' in run.text:
                            current_datetime = datetime.now()
                            current_date = current_datetime.strftime("%d-%b-%Y")
                            run.text = run.text.replace("<REPORT_DATE>", current_date)
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    # insert tool type image
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '<TOOL_IMAGE>' in p.text:
                        p.text = p.text.replace("<TOOL_IMAGE>", "")
                        new_p = OxmlElement("w:p")
                        p._p.addnext(new_p)
                        new_para = Paragraph(new_p, p._parent)
                        run = new_para.add_run()
                        if 'MDT' in FILENAME_PPTX and 'ORA' in FILENAME_PPTX:
                            run.add_picture('MDT.jpg', width=Inches(2.75), height=Inches(1))
                            run.add_text("\n")
                            run.add_picture('ORA.jpg', width=Inches(2.75), height=Inches(1))
                        elif 'MDT' in FILENAME_PPTX and XPT == 'XPT-':
                            run.add_picture('XPT.jpg', width=Inches(2.75), height=Inches(1))
                            run.add_text("\n")
                            run.add_picture('MDT.jpg', width=Inches(2.75), height=Inches(1))
                        elif 'ORA' in FILENAME_PPTX and XPT == 'XPT-':
                            run.add_picture('XPT.jpg', width=Inches(2.75), height=Inches(1))
                            run.add_text("\n")
                            run.add_picture('ORA.jpg', width=Inches(2.75), height=Inches(1))
                        elif 'MDT' in FILENAME_PPTX:
                            run.add_picture('MDT.jpg', width=Inches(2.75), height=Inches(1.2))
                        elif 'ORA' in FILENAME_PPTX:
                            run.add_picture('ORA.jpg', width=Inches(2.75), height=Inches(1.2))

                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # edit footer text
    for section in document.sections:
        footer = section.footer
        footer.paragraphs[0].runs[0].text = footer.paragraphs[0].runs[0].text.replace("<WELL_NAME>", WELL_NAME)
        footer.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    # add stations summary text in the report
    for p in document.paragraphs:
        if 'Key DFA/Sampling observations' in p.text:
            for station in reversed(stations_summary_text):
                new_p = OxmlElement("w:p")
                p._p.addnext(new_p)
                new_para = Paragraph(new_p, p._parent)
                new_para.add_run(station).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                new_para.style = 'Bullet Key Observations'

    # add sampling summary tables

    # list of sampling summary file names
    sampling_summaries = [f for f in glob.glob("*.png") if "Sampling_summary" in f]
    DFA_pictures = [f for f in glob.glob("*.png") if "DFA" in f]
    # print(DFA_pictures)

    for p in document.paragraphs:
        if ('Sampling Summary Table' in p.text) and p.style.name == 'Head2':
            for picture in reversed(sampling_summaries):
                new_p = OxmlElement("w:p")
                p._p.addnext(new_p)
                new_para = Paragraph(new_p, p._parent)
                run = new_para.add_run()
                run.add_picture(picture, width=IMAGE_WIDTH_HORIZONTAL)
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.add_break(WD_BREAK.PAGE)

    # insert texts/pictures for each DFA/sampling station

    for p in document.paragraphs:
        if ('Sampling Stations' in p.text) and p.style.name == 'Head1':
            for station in reversed(stations_data):
                new_p = OxmlElement("w:p")
                p._p.addnext(new_p)
                new_para = Paragraph(new_p, p._parent)

                # add station header
                station_header = station['station_type'] + ' Station at ' + station['depth'] + ' ft MD (File ' + \
                                 station['file_number'] + ')'
                new_para.add_run(station_header).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                new_para.style = 'Head2'

                # add 'IPTT' header if pumped some volume/for some time at the station
                if (station['time'] != '-' and station['volume'] != '-'):
                    IPTT_header_p = OxmlElement("w:p")
                    new_para._p.addnext(IPTT_header_p)
                    IPTT_header_para = Paragraph(IPTT_header_p, new_para._parent)
                    run = IPTT_header_para.add_run(
                        'Buildup Pressure Derivative Plot @ ' + station['depth'] + ' ft MD (File-' + station['file_number'] + ')')
                    IPTT_header_para.style = 'Head3'

                    # adding boilerplate IPTT text
                    IPTT_default_text_p = OxmlElement("w:p")
                    IPTT_header_para._p.addnext(IPTT_default_text_p)
                    IPTT_default_text_para = Paragraph(IPTT_default_text_p, IPTT_header_para._parent)
                    run_text = IPTT_default_text_para.add_run('Place your IPTT interpretation here')
                    run_text.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    IPTT_default_text_para.style = 'StationText'
                    run_text.add_break(WD_BREAK.PAGE)

                # add 'Contamination' header and images for DFA/sampling stations

                # checking actual file names of the pictures
                Contamination_image_exists = False

                if os.path.exists(station['file_number'] + '_' + str(float(station['depth'])) + '_Contamination.png'):
                    Contamination_image = station['file_number'] + '_' + str(float(station['depth'])) + '_Contamination.png'
                    Contamination_image_exists = True
                elif os.path.exists(station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_Contamination.png'):
                    Contamination_image = station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_Contamination.png'
                    Contamination_image_exists = True
                elif os.path.exists(station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_Contamination.png'):
                    Contamination_image = station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_Contamination.png'
                    Contamination_image_exists = True

                if Contamination_image_exists:
                    # Contamination_image = station['file_number'] + '_' + station['depth'] + '_Contamination.png'
                    Contamination_header_p = OxmlElement("w:p")
                    new_para._p.addnext(Contamination_header_p)
                    Contamination_header_para = Paragraph(Contamination_header_p, new_para._parent)
                    Contamination_header_para.add_run(
                        'Contamination Plot @ ' + station['depth'] + ' ft MD (File-' + station['file_number'] + ')')
                    Contamination_header_para.style = 'Head3'

                    # adding 'Contamination' image
                    Contamination_image_p = OxmlElement("w:p")
                    Contamination_header_para._p.addnext(Contamination_image_p)
                    Contamination_image_para = Paragraph(Contamination_image_p, Contamination_header_para._parent)
                    run = Contamination_image_para.add_run()
                    run.add_picture(Contamination_image, width=IMAGE_WIDTH_VERTICAL)
                    Contamination_image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.add_break(WD_BREAK.PAGE)

                # add 'HC Composition' header and images for DFA/sampling stations

                # checking actual file names of the pictures
                HCcomp_image_exists = False

                if os.path.exists(station['file_number'] + '_' + str(float(station['depth'])) + '_HCcomp.png'):
                    HCcomp_image = station['file_number'] + '_' + str(float(station['depth'])) + '_HCcomp.png'
                    HCcomp_table = station['file_number'] + '_' + str(float(station['depth'])) + '_HCcompTable.png'
                    HCcomp_image_exists = True
                elif os.path.exists(station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_HCcomp.png'):
                    HCcomp_image = station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_HCcomp.png'
                    HCcomp_table = station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_HCcompTable.png'
                    HCcomp_image_exists = True
                elif os.path.exists(station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_HCcomp.png'):
                    HCcomp_image = station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_HCcomp.png'
                    HCcomp_table = station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_HCcompTable.png'
                    HCcomp_image_exists = True

                if HCcomp_image_exists:
                    HCcomp_header_p = OxmlElement("w:p")
                    new_para._p.addnext(HCcomp_header_p)
                    HCcomp_header_para = Paragraph(HCcomp_header_p, new_para._parent)
                    if station['tool_type'] == 'MDT':
                        HCcomp_header_para.add_run('IFA HC composition @ ' + station['depth'] + ' ft MD (File-' + station['file_number'] + ')')
                    elif station['tool_type'] == 'ORA':
                        HCcomp_header_para.add_run('FISO HC composition @ ' + station['depth'] + ' ft MD (File-' + station['file_number'] + ')')
                    HCcomp_header_para.style = 'Head3'

                    # adding 'HCcompTable' image
                    HCcompTable_p = OxmlElement("w:p")
                    HCcomp_header_para._p.addnext(HCcompTable_p)
                    HCcompTable_para = Paragraph(HCcompTable_p, HCcomp_header_para._parent)
                    run = HCcompTable_para.add_run()
                    run.add_picture(HCcomp_table)
                    HCcompTable_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.add_break(WD_BREAK.PAGE)

                    # adding 'HCcomp' image
                    HCcomp_image_p = OxmlElement("w:p")
                    HCcomp_header_para._p.addnext(HCcomp_image_p)
                    HCcomp_image_para = Paragraph(HCcomp_image_p, HCcomp_header_para._parent)
                    run = HCcomp_image_para.add_run()
                    run.add_picture(HCcomp_image, width=IMAGE_WIDTH_VERTICAL)
                    HCcomp_image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER



                # add 'Sample capture' header and images for sampling stations
                if station['station_type'] == 'Sampling':
                    # checking actual file names of the pictures
                    if os.path.exists(
                            station['file_number'] + '_' + station['depth'] + '_Sample_capture_1.png'):
                            Sample_capture_pictures = [
                                f for f in glob.glob("*.png") \
                                if (station['file_number'] in f \
                                and station['depth'] in f \
                                and "Sample_capture" in f)]
                    elif os.path.exists(station['file_number'] + '_' + str(
                            float(math.ceil(float(station['depth'])))) + '_Sample_capture_1.png'):
                            Sample_capture_pictures = [
                                f for f in glob.glob("*.png") \
                                if (station['file_number'] in f \
                                    and str(float(math.ceil(float(station['depth'])))) in f \
                                    and "Sample_capture" in f)]
                    elif os.path.exists(station['file_number'] + '_' + str(
                            float(math.floor(float(station['depth'])))) + '_Sample_capture_1.png'):
                            Sample_capture_pictures = [
                                f for f in glob.glob("*.png") \
                                if (station['file_number'] in f \
                                    and str(float(math.floor(float(station['depth'])))) in f \
                                    and "Sample_capture" in f)]
                    else:
                        Sample_capture_pictures = []

                    # is sample capture plots exist - add sample capture header and pictures
                    if len(Sample_capture_pictures) != 0:
                        # print(Sample_capture_pictures)
                        Sample_header_p = OxmlElement("w:p")
                        new_para._p.addnext(Sample_header_p)
                        Sample_header_para = Paragraph(Sample_header_p, new_para._parent)
                        Sample_header_para.add_run('Sample Capture Plot @ ' + station['depth'] + ' ft MD (File-' + station['file_number'] + ')')
                        Sample_header_para.style = 'Head3'

                        # adding 'Sample capture plots'
                        for picture in reversed(Sample_capture_pictures):
                            sample_image_p = OxmlElement("w:p")
                            Sample_header_para._p.addnext(sample_image_p)
                            sample_image_para = Paragraph(sample_image_p, Sample_header_para._parent)
                            run = sample_image_para.add_run()
                            pic = run.add_picture(picture, width=IMAGE_WIDTH_VERTICAL)
                            # print(picture)
                            # print(pic)
                            # print('HEIGHT:' + str(pic.height.inches))
                            # print('WIDTH:' + str(pic.width.inches))
                            sample_image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run.add_break(WD_BREAK.PAGE)

                # add 'IFA Log' header and 'IFA Log' image

                # checking actual file names of the pictures
                IFA_Log_image_exists = False

                if os.path.exists(station['file_number'] + '_' + str(float(station['depth'])) + '_Complete_station.png'):
                    IFA_Log_image = station['file_number'] + '_' + str(float(station['depth'])) + '_Complete_station.png'
                    IFA_Log_image_exists = True
                elif os.path.exists(station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_Complete_station.png'):
                    IFA_Log_image = station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_Complete_station.png'
                    IFA_Log_image_exists = True
                elif os.path.exists(station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_Complete_station.png'):
                    IFA_Log_image = station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_Complete_station.png'
                    IFA_Log_image_exists = True

                if IFA_Log_image_exists:
                    IFA_log_header_p = OxmlElement("w:p")
                    new_para._p.addnext(IFA_log_header_p)
                    IFA_log_header_para = Paragraph(IFA_log_header_p, new_para._parent)
                    if station['tool_type'] == 'MDT':
                        IFA_log_header_para.add_run('IFA Plot @ ' + station['depth'] + ' ft MD (File-' + station['file_number'] + ')')
                    elif station['tool_type'] == 'ORA':
                        IFA_log_header_para.add_run('FISO Plot @ ' + station['depth'] + ' ft MD (File-' + station['file_number'] + ')')
                    IFA_log_header_para.style = 'Head3'

                    # adding 'IFA Log' image
                    IFA_log_image_p = OxmlElement("w:p")
                    IFA_log_header_para._p.addnext(IFA_log_image_p)
                    IFA_log_image_para = Paragraph(IFA_log_image_p, IFA_log_header_para._parent)
                    run = IFA_log_image_para.add_run()
                    run.add_picture(IFA_Log_image, width=IMAGE_WIDTH_VERTICAL)
                    IFA_log_image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.add_break(WD_BREAK.PAGE)

                # add DFA picture
                DFA_picture_p = OxmlElement("w:p")
                new_para._p.addnext(DFA_picture_p)
                DFA_picture_para = Paragraph(DFA_picture_p, new_para._parent)
                run = DFA_picture_para.add_run()
                if os.path.exists(station['file_number'] + '_' + str(float(station['depth'])) + '_DFA.png'):
                    DFA_filename = station['file_number'] + '_' + str(float(station['depth'])) + '_DFA.png'
                elif os.path.exists(station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_DFA.png'):
                    DFA_filename = station['file_number'] + '_' + str(float(math.ceil(float(station['depth'])))) + '_DFA.png'
                else:
                    DFA_filename = station['file_number'] + '_' + str(float(math.floor(float(station['depth'])))) + '_DFA.png'
                run.add_picture(DFA_filename, width=IMAGE_WIDTH_VERTICAL)
                DFA_picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                DFA_picture_para.style = 'DFAPicStyle'
                run.add_break(WD_BREAK.PAGE)

                # add station text
                station_text_p = OxmlElement("w:p")
                new_para._p.addnext(station_text_p)
                station_text_para = Paragraph(station_text_p, new_para._parent)
                station_text_para.add_run(station['text']).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                station_text_para.style = 'StationText'



    document.save(FILENAME_DOCX_REPORT + '.docx')

    progress_bar['value'] += 20
    progress_label['text'] = 'Generating docx report...[OK]'
    progress_bar.update_idletasks()

    # delete all png files in the current directory
    for f in glob.glob("*.png"):
        os.remove(f)

    progress_bar['value'] += 20
    progress_label['text'] = 'Deleting temp files...[OK]'
    progress_bar.update_idletasks()



